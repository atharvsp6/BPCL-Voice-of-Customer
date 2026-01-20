"""
Convert VoC_Hello_BPCL_Project_Report.md to Word Document (.docx)
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
import re

def create_styled_document():
    """Create a Word document with custom styles"""
    doc = Document()
    
    # Set up page margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    return doc

def parse_markdown_to_docx(md_file, docx_file):
    """Convert markdown file to Word document"""
    
    doc = create_styled_document()
    
    with open(md_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    i = 0
    in_code_block = False
    in_table = False
    table_data = []
    
    while i < len(lines):
        line = lines[i].rstrip()
        
        # Skip empty lines in most cases
        if not line.strip():
            i += 1
            continue
        
        # Handle code blocks
        if line.startswith('```'):
            in_code_block = not in_code_block
            i += 1
            continue
        
        if in_code_block:
            p = doc.add_paragraph(line, style='Normal')
            p.runs[0].font.name = 'Courier New'
            p.runs[0].font.size = Pt(9)
            i += 1
            continue
        
        # Handle headers
        if line.startswith('# '):
            p = doc.add_heading(line[2:], level=1)
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('#### '):
            doc.add_heading(line[5:], level=4)
        
        # Handle horizontal rules
        elif line.startswith('---'):
            p = doc.add_paragraph()
            p.add_run('_' * 100)
        
        # Handle tables
        elif '|' in line and not line.startswith('!'):
            if not in_table:
                in_table = True
                table_data = []
            
            # Parse table row
            cells = [cell.strip() for cell in line.split('|')[1:-1]]
            
            # Skip separator rows
            if all(set(cell.replace('-', '').strip()) <= {''} for cell in cells):
                i += 1
                continue
            
            table_data.append(cells)
            
            # Check if next line is still part of table
            if i + 1 < len(lines):
                next_line = lines[i + 1].strip()
                if '|' not in next_line or next_line.startswith('!'):
                    # End of table, create it
                    if table_data:
                        num_cols = max(len(row) for row in table_data)
                        table = doc.add_table(rows=len(table_data), cols=num_cols)
                        table.style = 'Light Grid Accent 1'
                        
                        for row_idx, row_data in enumerate(table_data):
                            for col_idx, cell_data in enumerate(row_data):
                                cell = table.rows[row_idx].cells[col_idx]
                                # Clean markdown formatting
                                cell_text = re.sub(r'\*\*([^*]+)\*\*', r'\1', cell_data)
                                cell.text = cell_text
                                
                                # Bold header row
                                if row_idx == 0:
                                    for paragraph in cell.paragraphs:
                                        for run in paragraph.runs:
                                            run.font.bold = True
                        
                        doc.add_paragraph()  # Add spacing after table
                    in_table = False
                    table_data = []
        
        # Handle lists
        elif line.startswith('- ') or line.startswith('* '):
            text = line[2:]
            # Remove markdown bold
            text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
            p = doc.add_paragraph(text, style='List Bullet')
        
        elif re.match(r'^\d+\. ', line):
            text = re.sub(r'^\d+\. ', '', line)
            text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
            p = doc.add_paragraph(text, style='List Number')
        
        # Handle images/figures (as text references)
        elif line.startswith('!['):
            match = re.match(r'!\[([^\]]+)\]\(([^)]+)\)', line)
            if match:
                alt_text, path = match.groups()
                p = doc.add_paragraph()
                run = p.add_run(f"[Figure: {path}]")
                run.italic = True
                run.font.color.rgb = RGBColor(128, 128, 128)
        
        # Handle math equations (convert to text)
        elif line.startswith('$$'):
            equation_lines = [line[2:]]
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('$$'):
                equation_lines.append(lines[i].rstrip())
                i += 1
            
            equation_text = '\n'.join(equation_lines)
            p = doc.add_paragraph()
            run = p.add_run(equation_text)
            run.font.name = 'Cambria Math'
            run.font.size = Pt(10)
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Handle inline math
        elif '$' in line:
            # Simple inline math handling
            text = re.sub(r'\$([^$]+)\$', r'[\1]', line)
            text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
            p = doc.add_paragraph(text)
        
        # Regular paragraphs
        else:
            # Clean markdown formatting
            text = re.sub(r'\*\*([^*]+)\*\*', r'\1', line)
            text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)  # Remove links but keep text
            
            if text.strip():
                p = doc.add_paragraph(text)
        
        i += 1
    
    # Save document
    doc.save(docx_file)
    print(f"✅ Successfully converted to: {docx_file}")

if __name__ == "__main__":
    md_file = "VoC_Hello_BPCL_Project_Report.md"
    docx_file = "VoC_Hello_BPCL_Project_Report.docx"
    
    print("Converting Markdown to Word Document...")
    parse_markdown_to_docx(md_file, docx_file)
    print("Done!")
